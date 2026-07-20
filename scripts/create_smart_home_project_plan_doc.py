from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/hoangminhan/face-recognition-door-lock")
OUT = ROOT / "outputs" / "ke-hoach-du-an-nha-thong-minh-ai.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "1F2937"
MUTED = "5B677A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="B8C2CC", sz="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color=DARK_BLUE, size=font_size)
        set_cell_shading(hdr[i], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    if widths:
        set_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    set_repeat_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    set_table_width(table, [6.5])
    set_table_borders(table, color="D6DEE8", sz="4")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("Kế Hoạch Chi Tiết Dự Án")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhà thông minh AI: khóa cửa nhận diện khuôn mặt, ổ cắm thông minh, điều khiển giọng nói và thang máy mô hình")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Phiên bản kế hoạch", "Bản mượt, triển khai bởi 1 người"],
            ["Mục tiêu", "Demo mượt hệ thống AI + IoT + smart home + dashboard realtime"],
            ["Thời gian khuyến nghị", "20 tuần cho 1 người làm chắc, có thời gian tích hợp và sửa lỗi"],
            ["Ngày lập", date.today().strftime("%d/%m/%Y")],
        ],
        widths=[1.8, 4.7],
        font_size=10,
    )

    add_callout(
        doc,
        "Định vị đề tài",
        "Dự án không chỉ là khóa cửa thông minh. Trọng tâm là một hệ thống nhà thông minh mini có AI nhận diện, điều phối thiết bị, lưu log, dashboard realtime và mô hình vật lý đủ trực quan để bảo vệ.",
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. Tóm tắt mục tiêu", 1)
    doc.add_paragraph(
        "Dự án xây dựng mô hình nhà thông minh tích hợp khóa cửa nhận diện khuôn mặt, điều khiển thiết bị tự động, ổ cắm thông minh độc lập, điều khiển giọng nói và thang máy mô hình. "
        "Hệ thống có backend, database, MQTT và web dashboard để quản lý người dùng, thiết bị, lịch sử ra vào và cảnh báo realtime."
    )
    add_bullets(
        doc,
        [
            "Tạo điểm khác biệt so với dự án IoT thuần phần cứng bằng lớp AI, backend, dashboard và quản trị dữ liệu.",
            "Đảm bảo demo có câu chuyện rõ: chủ nhà về, cửa mở, nhà tự bật thiết bị, có thể gọi thang máy bằng giọng nói.",
            "Giữ phạm vi có thể hoàn thiện: bản lõi phải chạy chắc trước, tính năng nâng cao thêm sau.",
        ],
    )

    add_heading(doc, "2. Phạm vi sản phẩm", 1)
    add_table(
        doc,
        ["Mức", "Tính năng", "Mục tiêu hoàn thiện"],
        [
            ["Core", "Nhận diện khuôn mặt, mở khóa, ghi log, dashboard trạng thái cửa", "Bắt buộc hoàn thành"],
            ["Smart Home", "Đèn, quạt, TV mô hình, ổ cắm thông minh DC, tự động theo ánh sáng/nhiệt độ", "Bắt buộc nếu muốn dự án nổi bật"],
            ["Voice", "Ra lệnh bằng giọng nói từ dashboard: mở cửa, bật/tắt thiết bị, gọi thang máy", "Nên hoàn thành"],
            ["Elevator", "Thang máy mô hình 2 tầng, có giới hạn hành trình và báo tầng", "Tính năng wow nếu còn thời gian"],
            ["Security", "Người lạ, cửa mở lâu, guest mode, cảnh báo Telegram", "Nâng cao nhưng rất đáng điểm"],
            ["Emergency", "Gas/flame/nhiệt độ bất thường -> mở cửa, còi, quạt hút", "Nâng cao"],
        ],
        widths=[1.0, 3.5, 2.0],
    )

    add_heading(doc, "3. Kiến trúc hệ thống", 1)
    doc.add_paragraph(
        "Kiến trúc đề xuất dùng nhiều node IoT giao tiếp qua MQTT. Backend là trung tâm điều phối, database lưu trạng thái/log, dashboard hiển thị realtime và nhận lệnh điều khiển."
    )
    add_table(
        doc,
        ["Thành phần", "Vai trò", "Giao tiếp chính", "Ghi chú"],
        [
            ["Node 1: Smart Door", "ESP32-CAM chụp ảnh, nhận lệnh mở khóa, điều khiển khóa/còi", "Wi-Fi, HTTP/MQTT", "Ưu tiên ổn định nguồn và tốc độ phản hồi"],
            ["Node 2: Smart Power Box", "Ổ cắm thông minh DC cho đèn/quạt/TV mô hình", "MQTT", "Dùng relay/MOSFET 4 kênh, không phụ thuộc máy tính cấp điện"],
            ["Node 3: Elevator", "Điều khiển thang máy 2 tầng, limit switch, LED tầng", "MQTT hoặc Serial/Wi-Fi", "Làm 2 tầng để giảm rủi ro cơ khí"],
            ["Node 4: Environment/Safety", "DHT22, LDR, gas/flame, quạt hút, còi", "MQTT", "Có thể gộp với Node 2 nếu thiếu thời gian"],
            ["Backend", "API, rule engine, auth, MQTT broker client, log", "REST, WebSocket, MQTT", "Node.js/Express theo repo hiện tại"],
            ["Dashboard", "Quản lý user, log, thiết bị, voice command", "REST, WebSocket", "Đây là lợi thế lớn so với IoT thuần"],
            ["AI Service", "Detect/recognize face, trả kết quả định danh", "HTTP/file/image stream", "Có thể chạy trên laptop/server thay vì ESP32-CAM"],
        ],
        widths=[1.45, 2.2, 1.25, 1.6],
    )

    add_heading(doc, "4. Luồng demo chuẩn", 1)
    add_numbers(
        doc,
        [
            "Chủ nhà đứng trước cửa, ESP32-CAM chụp ảnh và gửi về AI/backend.",
            "AI nhận diện Owner, backend kiểm tra quyền và ghi sự kiện.",
            "Backend gửi lệnh mở khóa qua MQTT/API tới Node Smart Door.",
            "Node Smart Power Box bật đèn nếu LDR báo tối, bật quạt nếu DHT22 báo nóng.",
            "Dashboard cập nhật realtime: cửa mở, Owner entered, đèn/quạt đang bật.",
            "Người dùng nói trên dashboard: gọi thang máy tầng 2.",
            "Node Elevator nhận lệnh, chạy motor tới tầng 2, limit switch xác nhận vị trí.",
            "Người lạ thử vào: cửa không mở, buzzer kêu, snapshot/log cảnh báo hiển thị trên dashboard.",
        ],
    )

    add_heading(doc, "5. Kế hoạch triển khai bản mượt 20 tuần cho 1 người", 1)
    add_table(
        doc,
        ["Tuần", "Mục tiêu", "Công việc chi tiết", "Kết quả bàn giao"],
        [
            ["1", "Chốt thiết kế", "Vẽ kiến trúc, chốt node, mua linh kiện, chuẩn hóa topic MQTT, chuẩn bị repo", "Sơ đồ hệ thống, BOM, backlog"],
            ["2", "Thiết kế chi tiết", "Vẽ sơ đồ node, sơ đồ nguồn, sơ đồ dây sơ bộ, chia topic MQTT/API, lập checklist test", "Thiết kế kỹ thuật v1"],
            ["3", "Backend nền", "Hoàn thiện API health/devices/users/access logs, cấu hình env, database schema", "Backend chạy local/Docker"],
            ["4", "MQTT nền", "Cài broker, publish/subscribe command/status, mock device test", "Backend gửi/nhận MQTT ổn"],
            ["5", "AI nhận diện v1", "Chuẩn bị dataset owner/family/guest, test face detector, embedding, threshold", "API nhận diện trả identity"],
            ["6", "AI nhận diện v2", "Test nhiều góc mặt/ánh sáng, xử lý unknown, lưu snapshot sự kiện", "Nhận diện ổn định hơn"],
            ["7", "Smart Door phần cứng", "ESP32-CAM, servo/relay, buzzer, test mở khóa độc lập", "Cửa mô hình mở/đóng ổn"],
            ["8", "Tích hợp cửa", "AI -> backend -> MQTT -> ESP32, lưu log, xử lý unknown", "Demo nhận diện mở cửa"],
            ["9", "Dashboard v1", "Trang trạng thái cửa, log ra vào, điều khiển thủ công", "Dashboard xem log/trạng thái"],
            ["10", "Dashboard realtime", "WebSocket/MQTT status, trạng thái thiết bị realtime, snapshot người quét", "Dashboard realtime"],
            ["11", "Smart socket v1", "ESP32 + relay/MOSFET 4 kênh, cổng DC cho đèn/quạt/TV", "Thiết bị cắm vào tự chạy"],
            ["12", "Smart socket v2", "Đóng hộp, LED trạng thái, nút tay fallback, test tải DC", "Ổ cắm thông minh ổn định"],
            ["13", "Automation", "Rule owner/guest/unknown, LDR bật đèn, DHT22 bật quạt", "Nhà tự phản ứng theo ngữ cảnh"],
            ["14", "Voice command", "Web Speech API hoặc module voice, parser lệnh, gửi command backend", "Bật/tắt thiết bị bằng giọng nói"],
            ["15", "Elevator cơ khí", "Khung 2 tầng, cabin, stepper/servo, limit switch, test lên/xuống", "Thang máy chạy bằng nút/test command"],
            ["16", "Elevator tích hợp", "MQTT command, dashboard trạng thái tầng, voice gọi tầng", "Gọi thang máy bằng giọng nói"],
            ["17", "Security/Emergency", "Cửa mở lâu, unknown nhiều lần, gas/flame cơ bản, Telegram nếu kịp", "Cảnh báo hoạt động"],
            ["18", "Hardening", "Test mất mạng, mất MQTT, nguồn yếu, delay AI, retry command, watchdog", "Demo chạy 5-7 phút ổn"],
            ["19", "Hoàn thiện mô hình", "Đi dây gọn, dán nhãn node, cố định nhà/thang máy/ổ cắm, tối ưu ánh sáng camera", "Mô hình đẹp và dễ bảo trì"],
            ["20", "Hoàn thiện nộp", "Viết báo cáo, slide, quay video, luyện demo, chuẩn bị phương án fallback", "Bộ mô hình + tài liệu hoàn chỉnh"],
        ],
        widths=[0.55, 1.35, 3.35, 1.25],
        font_size=7.8,
    )

    add_heading(doc, "6. Mốc kiểm soát cho 1 người", 1)
    add_table(
        doc,
        ["Mốc", "Tuần", "Điều kiện phải đạt"],
        [
            ["MVP cửa thông minh", "8", "Nhận diện Owner/Unknown, mở khóa, ghi log cơ bản"],
            ["Dashboard usable", "10", "Xem log, trạng thái cửa realtime, có nút điều khiển thủ công"],
            ["Smart home chạy độc lập", "13", "Đèn/quạt/TV mô hình bật bằng smart socket, không phụ thuộc máy tính cấp điện"],
            ["Voice + elevator", "16", "Ra lệnh giọng nói và gọi thang máy 2 tầng ở mức ổn định"],
            ["Bản demo khóa phạm vi", "18", "Không thêm tính năng lớn nữa, chỉ sửa lỗi và tối ưu"],
            ["Bản nộp", "20", "Mô hình gọn, tài liệu đủ, kịch bản demo đã luyện nhiều lần"],
        ],
        widths=[1.65, 0.75, 4.1],
    )

    add_heading(doc, "7. Kế hoạch rà soát kỹ và kiểm thử", 1)
    add_heading(doc, "7.1 Checklist theo cổng chất lượng", 2)
    add_table(
        doc,
        ["Cổng", "Điều kiện đạt", "Cách kiểm tra", "Không đạt thì xử lý"],
        [
            ["Gate 1: Backend", "API chạy, database ghi log, MQTT publish được", "Curl API, mqtt_sub/mqtt_pub, xem database", "Sửa schema/env/topic trước khi làm UI"],
            ["Gate 2: AI", "Nhận diện Owner ổn, Unknown không mở cửa", "Test ít nhất 30 ảnh/snapshot trong nhiều ánh sáng", "Điều chỉnh threshold/dataset/ánh sáng"],
            ["Gate 3: Door", "Servo/relay mở khóa đúng thời gian, không giật", "Test 30 lần mở/đóng liên tục", "Tách nguồn, thêm delay/debounce"],
            ["Gate 4: Dashboard", "Log/trạng thái realtime, nút điều khiển không lỗi", "Mở dashboard khi node gửi status", "Sửa WebSocket/state sync"],
            ["Gate 5: Socket", "Đèn/quạt/TV chạy độc lập khi cắm nguồn", "Bật/tắt từng kênh 20 lần", "Kiểm tra relay/MOSFET, nguồn, GND"],
            ["Gate 6: Voice", "Lệnh phổ biến nhận đúng trên môi trường demo", "Đọc 20 lệnh mẫu", "Giới hạn câu lệnh, thêm nút fallback"],
            ["Gate 7: Elevator", "Cabin dừng đúng tầng, không vượt hành trình", "Test lên/xuống 20 chu kỳ", "Hiệu chỉnh limit switch, tốc độ, dây kéo"],
            ["Gate 8: Full Demo", "Chạy kịch bản 5-7 phút không lỗi nghiêm trọng", "Tổng duyệt ít nhất 5 lần", "Cắt tính năng phụ nếu làm giảm độ ổn định"],
        ],
        widths=[1.25, 2.0, 1.65, 1.6],
        font_size=7.8,
    )

    add_heading(doc, "7.2 Test case bắt buộc", 2)
    add_table(
        doc,
        ["ID", "Tình huống", "Kết quả mong đợi", "Ưu tiên"],
        [
            ["TC-01", "Owner đứng trước cửa đủ sáng", "Nhận diện đúng, mở khóa, ghi log", "P0"],
            ["TC-02", "Owner đứng trong ánh sáng yếu", "Nhận diện hoặc yêu cầu chụp lại, không mở sai", "P0"],
            ["TC-03", "Unknown đứng trước cửa", "Không mở khóa, buzzer/log cảnh báo", "P0"],
            ["TC-04", "Guest còn hạn quyền", "Mở cửa, chỉ bật thiết bị cơ bản", "P1"],
            ["TC-05", "Guest hết hạn quyền", "Không mở, log access denied", "P1"],
            ["TC-06", "Cửa mở quá lâu", "Cảnh báo dashboard/buzzer", "P1"],
            ["TC-07", "Owner vào khi trời tối", "Đèn bật tự động", "P0"],
            ["TC-08", "Nhiệt độ cao", "Quạt bật tự động", "P1"],
            ["TC-09", "Nói 'bật đèn tầng 1'", "Đèn tầng 1 bật, status cập nhật", "P1"],
            ["TC-10", "Nói 'gọi thang máy tầng 2'", "Thang máy đến tầng 2, LED/status đúng", "P1"],
            ["TC-11", "Mất MQTT tạm thời", "Dashboard báo offline, không treo firmware", "P1"],
            ["TC-12", "Flame/gas kích hoạt", "Emergency mode, còi/quạt/cửa theo cấu hình", "P2"],
        ],
        widths=[0.7, 2.45, 2.65, 0.7],
        font_size=7.8,
    )

    add_heading(doc, "7.3 Rà soát kỹ thuật hằng tuần", 2)
    add_bullets(
        doc,
        [
            "Cuối mỗi tuần phải có video test ngắn hoặc ảnh chụp trạng thái để chứng minh module thật sự chạy.",
            "Không để module nào chỉ chạy riêng lẻ quá 1 tuần mà chưa nối vào MQTT/backend.",
            "Mỗi thay đổi topic MQTT/API phải cập nhật vào tài liệu interface để tránh lệch giữa firmware và backend.",
            "Luôn có phương án fallback bằng nút dashboard hoặc nút vật lý nếu voice/AI lỗi trong lúc demo.",
            "Trước tuần 18 phải khóa phạm vi; sau mốc này chỉ sửa lỗi, tối ưu và làm báo cáo.",
        ],
    )

    add_heading(doc, "8. Rủi ro và cách giảm rủi ro", 1)
    add_table(
        doc,
        ["Rủi ro", "Mức", "Dấu hiệu", "Cách giảm"],
        [
            ["Nhận diện mặt sai hoặc chậm", "Cao", "Mở nhầm/không mở, delay dài", "Giới hạn góc camera, thêm đèn phụ, tuning threshold, fallback RFID/keypad"],
            ["Nguồn ESP32-CAM yếu", "Cao", "Reset, mất Wi-Fi, camera lỗi", "Dùng nguồn 5V 2A, dây ngắn, tách nguồn motor/relay"],
            ["Relay/ổ cắm nguy hiểm nếu dùng 220V", "Cao", "Chập/cháy/giật", "Demo bằng DC 5V/12V; nếu AC phải hộp kín, cầu chì, người có kinh nghiệm kiểm tra"],
            ["Thang máy lệch tầng", "Trung bình", "Cabin không dừng đúng vị trí", "Dùng limit switch, tốc độ thấp, test cơ khí sớm"],
            ["Voice nhận sai lệnh", "Trung bình", "Bật nhầm thiết bị", "Dùng tập lệnh ngắn, xác nhận lệnh trên dashboard, có nút thủ công"],
            ["Tích hợp trễ deadline", "Cao", "Module rời rạc đến tuần cuối", "Tích hợp từng phần từ tuần 3-4, demo nhỏ mỗi tuần"],
            ["Dashboard chỉ để cho có", "Trung bình", "Không realtime, không log rõ", "Ưu tiên trạng thái/log/snapshot trước giao diện trang trí"],
        ],
        widths=[1.75, 0.75, 1.75, 2.25],
        font_size=7.8,
    )

    add_heading(doc, "9. Thống kê linh kiện cần mua", 1)
    add_callout(
        doc,
        "Nguyên tắc mua linh kiện",
        "Ưu tiên linh kiện dễ thay thế, dễ tìm, chạy 5V/12V cho mô hình. Không dùng điện 220V trong bản demo đầu tiên nếu không có hộp bảo vệ và người kiểm tra an toàn điện.",
        fill="FFF7D6",
    )

    add_heading(doc, "9.1 Linh kiện bắt buộc cho bản mượt", 2)
    add_table(
        doc,
        ["Nhóm", "Linh kiện", "SL", "Thông số gợi ý", "Dùng cho"],
        [
            ["Điều khiển", "ESP32-CAM AI-Thinker + camera OV2640", "1", "Cấp 5V ổn định, có Wi-Fi, camera", "Nhận diện/cửa"],
            ["Điều khiển", "ESP32 DevKit / ESP8266 NodeMCU", "2-3", "Wi-Fi 2.4 GHz, đủ GPIO", "Smart socket, elevator, môi trường"],
            ["Nạp/debug", "USB to TTL CP2102/CH340", "1", "Hỗ trợ 5V/3.3V", "Nạp ESP32-CAM"],
            ["Khóa cửa", "Servo MG996R hoặc solenoid lock 12V", "1", "Servo cho mô hình; solenoid cho demo khóa thật", "Cửa chính"],
            ["Công suất", "Relay module 4 kênh", "1", "5V, opto cách ly nếu có", "Đèn/quạt/TV/khóa"],
            ["Công suất", "MOSFET driver module 4 kênh", "1", "Cho tải DC, chạy 5V/12V", "Ổ cắm DC thay relay nếu cần êm hơn"],
            ["Nguồn", "Adapter 12V DC", "1", "Tối thiểu 2A, ưu tiên 3A", "Solenoid, quạt, LED 12V"],
            ["Nguồn", "Buck converter LM2596/MP1584", "2", "12V xuống 5V, 2A trở lên", "ESP32, relay, motor"],
            ["Mô hình", "Breadboard/board hàn lỗ", "2-3", "Loại vừa", "Lắp mạch"],
            ["Mô hình", "Dây jumper + dây điện chịu tải", "1 bộ", "Dupont và dây nguồn lớn hơn", "Đấu nối"],
            ["Kết nối", "Terminal block/jack DC", "6-10", "2 pin/3 pin", "Nguồn và tải"],
            ["Hiển thị", "OLED 0.96 inch I2C", "1", "SSD1306, 3.3V/5V", "Trạng thái tầng/thiết bị"],
            ["Cảm biến", "DHT22/AM2302", "1", "3.3-5V, nhiệt/ẩm", "Bật quạt theo nhiệt độ"],
            ["Cảm biến", "LDR + điện trở 10k", "1", "Đọc ánh sáng analog", "Bật đèn khi tối"],
        ],
        widths=[0.9, 2.0, 0.45, 2.0, 1.15],
        font_size=7.4,
    )

    add_heading(doc, "9.2 Linh kiện thang máy mô hình", 2)
    add_table(
        doc,
        ["Linh kiện", "SL", "Thông số gợi ý", "Ghi chú"],
        [
            ["Stepper motor 28BYJ-48 + driver ULN2003", "1", "5V, geared stepper", "Dễ điều khiển tầng, phù hợp mô hình nhẹ"],
            ["Limit switch / công tắc hành trình", "2-3", "Một cái mỗi tầng hoặc đầu hành trình", "Bắt buộc để chống vượt tầng"],
            ["Dây cước/dây kéo + pulley nhỏ", "1 bộ", "Nhẹ, ít ma sát", "Kéo cabin"],
            ["Ray trượt mica/gỗ/nhôm", "1 bộ", "Cố định cabin thẳng", "Giảm kẹt cơ khí"],
            ["LED báo tầng", "2-3", "Kèm điện trở 220 ohm", "Hiển thị tầng hiện tại"],
            ["Nút nhấn gọi tầng", "2", "Nút vật lý fallback", "Dùng khi voice lỗi"],
        ],
        widths=[2.0, 0.5, 2.0, 2.0],
        font_size=7.8,
    )

    add_heading(doc, "9.3 Linh kiện an ninh và khẩn cấp", 2)
    add_table(
        doc,
        ["Linh kiện", "SL", "Thông số gợi ý", "Mức ưu tiên"],
        [
            ["Cảm biến cửa từ", "1", "Reed switch", "P1"],
            ["Buzzer/còi 5V", "1-2", "Cảnh báo người lạ/cửa mở lâu", "P1"],
            ["Keypad 4x4 hoặc RFID RC522", "1", "Mở khóa dự phòng", "P2"],
            ["MQ-2 hoặc MQ-135", "1", "Gas/khói cơ bản", "P2"],
            ["Flame sensor", "1", "Phát hiện lửa trong demo", "P2"],
            ["Quạt hút mini 5V/12V", "1", "Emergency/ventilation", "P2"],
            ["PIR hoặc mmWave", "1", "Phát hiện có người trong phòng", "P3"],
        ],
        widths=[2.0, 0.5, 2.9, 1.1],
        font_size=7.8,
    )

    add_heading(doc, "9.4 Vật tư mô hình và dụng cụ", 2)
    add_table(
        doc,
        ["Vật tư/dụng cụ", "SL", "Ghi chú"],
        [
            ["Formex/mica/gỗ làm nhà 2 tầng", "1 bộ", "Nên thiết kế tháo lắp được để sửa dây"],
            ["Ke góc, ốc vít, keo nến, băng keo 2 mặt", "1 bộ", "Cố định node và dây"],
            ["Hộp nhựa cách điện cho smart socket", "1-2", "Rất quan trọng nếu có phần công suất"],
            ["Đồng hồ đo điện", "1", "Đo 5V/12V, kiểm tra chập"],
            ["Mỏ hàn + thiếc + ống co nhiệt", "1 bộ", "Làm dây chắc, tránh lỏng khi demo"],
            ["Nhãn dán dây/node", "1 bộ", "Giúp bảo trì và thuyết trình dễ hiểu"],
        ],
        widths=[2.5, 0.6, 3.4],
        font_size=7.8,
    )

    add_heading(doc, "9.5 Tổng hợp số lượng theo nhóm", 2)
    add_table(
        doc,
        ["Nhóm", "Số món ước tính", "Mức cần thiết"],
        [
            ["Core door + AI hardware", "5-6", "Bắt buộc"],
            ["Smart socket/smart home", "7-9", "Bắt buộc để nổi bật"],
            ["Elevator", "6", "Nên có nếu còn thời gian"],
            ["Security/Emergency", "4-7", "Nâng cao"],
            ["Vật tư mô hình/dụng cụ", "6", "Bắt buộc để demo đẹp"],
            ["Tổng mua tối thiểu bản mượt", "24-28", "Không tính linh kiện dự phòng nhỏ"],
            ["Tổng nếu mua đủ nâng cao", "34-40", "Có dự phòng và emergency"],
        ],
        widths=[2.4, 1.4, 2.7],
        font_size=8,
    )

    add_heading(doc, "10. Dự toán giá thành bản mượt", 1)
    add_callout(
        doc,
        "Cách đọc bảng giá",
        "Giá là khoảng ước tính tại thời điểm lập kế hoạch, tính bằng VND và chưa gồm phí vận chuyển. Khi mua thật nên cộng thêm 10-15% dự phòng vì giá shop, chất lượng module và tồn kho thay đổi.",
        fill="FFF7D6",
    )
    add_table(
        doc,
        ["Nhóm", "Linh kiện/gói mua", "SL", "Giá thấp", "Giá cao", "Thành tiền ước tính"],
        [
            ["Core", "ESP32-CAM AI-Thinker", "1", "180.000", "220.000", "180.000-220.000"],
            ["Core", "ESP32 DevKit/ESP8266 NodeMCU", "3", "120.000", "172.000", "360.000-516.000"],
            ["Core", "USB to TTL CP2102/CH340", "1", "40.000", "80.000", "40.000-80.000"],
            ["Door", "Servo MG996R hoặc solenoid lock 12V", "1", "60.000", "180.000", "60.000-180.000"],
            ["Power", "Relay 4 kênh 5V", "1", "44.000", "75.000", "44.000-75.000"],
            ["Power", "MOSFET driver 4 kênh", "1", "50.000", "120.000", "50.000-120.000"],
            ["Power", "Adapter 12V 2-3A", "1", "80.000", "150.000", "80.000-150.000"],
            ["Power", "Buck converter LM2596/MP1584", "2", "15.000", "40.000", "30.000-80.000"],
            ["Wiring", "Breadboard/board hàn lỗ", "2-3", "60.000", "150.000", "60.000-150.000"],
            ["Wiring", "Dây jumper + dây điện chịu tải", "1 bộ", "80.000", "150.000", "80.000-150.000"],
            ["Wiring", "Terminal block/jack DC", "1 bộ", "50.000", "120.000", "50.000-120.000"],
            ["Display", "OLED 0.96 inch I2C", "1", "50.000", "90.000", "50.000-90.000"],
            ["Sensor", "DHT22/AM2302", "1", "77.000", "140.000", "77.000-140.000"],
            ["Sensor", "LDR + điện trở", "1 bộ", "10.000", "25.000", "10.000-25.000"],
            ["Elevator", "28BYJ-48 + ULN2003", "1", "35.000", "70.000", "35.000-70.000"],
            ["Elevator", "Limit switch", "3", "5.000", "15.000", "15.000-45.000"],
            ["Elevator", "Dây kéo/pulley/ray/cabin/LED/nút", "1 bộ", "170.000", "450.000", "170.000-450.000"],
            ["Security", "Cảm biến cửa từ + buzzer", "1 bộ", "25.000", "80.000", "25.000-80.000"],
            ["Emergency", "MQ-2/MQ-135 + flame sensor + quạt mini", "1 bộ", "120.000", "300.000", "120.000-300.000"],
            ["Model", "Formex/mica/gỗ, ke, ốc, hộp nhựa", "1 bộ", "250.000", "600.000", "250.000-600.000"],
            ["Tooling", "Thiếc, ống co nhiệt, nhãn dây, vật tư hao mòn", "1 bộ", "150.000", "400.000", "150.000-400.000"],
        ],
        widths=[0.85, 2.15, 0.45, 0.85, 0.85, 1.35],
        font_size=6.8,
    )
    add_table(
        doc,
        ["Khoản", "Ước tính"],
        [
            ["Tổng thấp bản mượt", "Khoảng 1.9 triệu VND"],
            ["Tổng cao bản mượt", "Khoảng 4.0 triệu VND"],
            ["Dự phòng 15%", "Khoảng 0.3-0.6 triệu VND"],
            ["Ngân sách nên chuẩn bị", "Khoảng 2.2-4.6 triệu VND"],
            ["Không tính trong bảng", "Laptop/server chạy AI, phí ship, dụng cụ lớn nếu đã có như mỏ hàn/đồng hồ đo"],
        ],
        widths=[2.25, 4.25],
        font_size=8,
    )

    add_heading(doc, "11. Chuẩn giao tiếp MQTT/API", 1)
    add_table(
        doc,
        ["Topic/API", "Chiều", "Payload mẫu", "Mục đích"],
        [
            ["doorlock/device/{id}/command", "Backend -> Node", '{"action":"unlock","durationMs":5000}', "Mở/khóa cửa"],
            ["doorlock/device/{id}/status", "Node -> Backend", '{"lock":"locked","online":true}', "Báo trạng thái node"],
            ["smarthome/socket/{id}/command", "Backend -> Node", '{"channel":"fan","action":"on"}', "Bật/tắt ổ cắm"],
            ["smarthome/elevator/{id}/command", "Backend -> Node", '{"targetFloor":2}', "Gọi thang máy"],
            ["smarthome/sensor/{id}/status", "Node -> Backend", '{"temp":31,"light":210}', "Dữ liệu cảm biến"],
            ["/api/recognition/events", "AI -> Backend", '{"userId":"owner_001","decision":"unlock"}', "Gửi kết quả nhận diện"],
        ],
        widths=[1.9, 1.15, 2.25, 1.2],
        font_size=7.4,
    )

    add_heading(doc, "12. Tiêu chí chấm điểm nội bộ", 1)
    add_table(
        doc,
        ["Hạng mục", "Điểm mục tiêu", "Điều kiện đạt điểm cao"],
        [
            ["Ý tưởng và tính ứng dụng", "9/10", "Có câu chuyện smart home rõ, không chỉ bật/tắt đơn giản"],
            ["AI nhận diện", "8.5-9/10", "Owner/Guest/Unknown hoạt động ổn, có log ảnh/sự kiện"],
            ["IoT/nhúng", "9/10", "Nhiều node chạy độc lập, MQTT ổn, thiết bị không phụ thuộc máy tính cấp điện"],
            ["Dashboard/backend", "9/10", "Realtime, quản lý user, log, điều khiển thủ công, trạng thái thiết bị"],
            ["Mô hình vật lý", "9/10", "Nhà 2 tầng, cửa, ổ cắm, thang máy, dây gọn, có nhãn"],
            ["Độ ổn định demo", "9/10", "Chạy 5-7 phút nhiều lần không lỗi lớn"],
            ["Tài liệu/báo cáo", "8.5-9/10", "Có kiến trúc, BOM, sơ đồ đấu dây, test case, rủi ro"],
        ],
        widths=[2.0, 1.1, 3.4],
        font_size=8,
    )

    add_heading(doc, "13. Nguồn tham khảo thông số linh kiện và giá", 1)
    add_bullets(
        doc,
        [
            "Giá ESP32 DevKit tham khảo từ Shopee/Linh Kiện Việt/Pyworld, dao động khoảng 120.000-172.000 VND tùy bản và shop.",
            "Giá relay 4 kênh tham khảo từ ThegioiIC, Điện tử 360, Hshop và Điện tử Tuyết Nga, khoảng 44.000-75.000 VND.",
            "Giá ESP32-CAM tham khảo từ listing Shopee khoảng 220.000 VND; khi mua có thể chọn shop khác rẻ hơn nếu đủ camera OV2640 và board nạp phù hợp.",
            "Giá DHT22 tham khảo từ Shopee/Lazada/Cytron/Nshop, khoảng 77.000-139.000 VND tùy dạng module và tồn kho.",
            "ESP32-CAM: tài liệu DFRobot/DigiKey và Handson Technology ghi nhận module dùng ESP32-S, camera OV2640, cấp 5V và có chế độ sleep.",
            "DHT22/AM2302: SparkFun/Adafruit mô tả nguồn 3.3-6V hoặc 3-5V, đo ẩm 0-100%RH và nhiệt độ -40 đến 80°C.",
            "DS3231: Analog Devices mô tả RTC I2C chính xác cao, hỗ trợ bus I2C và bù năm nhuận tới 2100.",
            "28BYJ-48: datasheet và hướng dẫn DigiKey/Mouser ghi motor stepper 5V, 4 pha, tỉ số giảm tốc khoảng 64:1.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "14. Câu chuyện demo bằng giấy: Khách và Admin", 1)
    doc.add_paragraph(
        "Phần này dùng để diễn giải hệ thống bằng giấy khi thuyết trình, hoặc dùng làm kịch bản dự phòng nếu một module phần cứng chưa kịp chạy. "
        "Mục tiêu là giúp người xem hiểu rõ vai trò Admin, Khách, phân quyền truy cập và cách hệ thống ra quyết định."
    )
    add_callout(
        doc,
        "Thông điệp chính khi demo",
        "Admin không chỉ mở cửa. Admin quản lý quyền vào nhà, cấp quyền cho khách, xem log và kiểm soát thiết bị. Khách chỉ được vào trong phạm vi được cho phép, không có quyền chỉnh cấu hình hệ thống.",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["Đạo cụ giấy", "Nội dung cần in/viết", "Cách dùng khi demo"],
        [
            ["Thẻ Admin", "Tên: Chủ nhà/Admin; Quyền: Full access", "Đưa trước camera hoặc giơ lên khi giới thiệu người có quyền cao nhất"],
            ["Thẻ Khách", "Tên: Guest A; Quyền: vào từ 08:00-18:00; chỉ bật đèn lối vào", "Dùng để minh họa khách hợp lệ nhưng quyền bị giới hạn"],
            ["Thẻ Unknown", "Không có tên trong hệ thống", "Dùng cho tình huống người lạ, hệ thống từ chối và cảnh báo"],
            ["Phiếu quyền truy cập", "Owner, Family, Guest, Unknown, Blocked", "Giải thích bảng phân quyền trước khi chạy demo"],
            ["Ảnh giấy khuôn mặt", "Ảnh Admin, ảnh Khách, ảnh Người lạ", "Mô phỏng dữ liệu camera/AI nếu cần trình bày offline"],
            ["Phiếu log mẫu", "10:05 Admin entered; 10:10 Guest entered; 10:20 Unknown denied", "Cho thấy hệ thống có lưu vết, không chỉ điều khiển phần cứng"],
            ["Sơ đồ nhà giấy", "Cửa, phòng khách, đèn, quạt, thang máy tầng 1-2", "Chỉ vào từng vùng khi giải thích hành động tự động"],
        ],
        widths=[1.5, 2.4, 2.6],
        font_size=7.8,
    )
    add_heading(doc, "14.1 Kịch bản lời dẫn 5 phút", 2)
    add_table(
        doc,
        ["Mốc", "Người diễn", "Hành động", "Lời dẫn ngắn"],
        [
            ["1", "Admin", "Giơ thẻ Admin và ảnh khuôn mặt Admin", "Đây là chủ nhà, có toàn quyền mở cửa, quản lý khách và điều khiển thiết bị."],
            ["2", "Hệ thống", "Chỉ vào sơ đồ camera -> backend -> dashboard", "Camera gửi ảnh về AI. Backend kiểm tra quyền và ghi log trước khi ra lệnh mở cửa."],
            ["3", "Admin", "Đặt ảnh Admin trước camera hoặc đưa thẻ lên", "Khi Admin được nhận diện, cửa mở, đèn bật nếu tối, quạt bật nếu nhiệt độ cao."],
            ["4", "Khách", "Giơ thẻ Khách", "Khách này đã được Admin cấp quyền, nhưng chỉ được vào trong khung giờ cho phép."],
            ["5", "Hệ thống", "Chỉ phiếu quyền truy cập của Khách", "Hệ thống mở cửa cho khách hợp lệ, nhưng chỉ bật đèn lối vào và không cho chỉnh cấu hình."],
            ["6", "Unknown", "Giơ thẻ Unknown", "Người lạ không có trong danh sách nên cửa không mở, buzzer cảnh báo và dashboard ghi log từ chối."],
            ["7", "Admin", "Chỉ phiếu log mẫu", "Admin xem được toàn bộ lịch sử: ai vào, lúc nào, kết quả nhận diện và thiết bị đã bật/tắt."],
            ["8", "Admin", "Nói lệnh giả lập: gọi thang máy tầng 2", "Lệnh giọng nói đi qua dashboard/backend, sau đó node thang máy nhận lệnh và di chuyển."],
        ],
        widths=[0.5, 1.0, 2.05, 2.95],
        font_size=7.6,
    )
    add_heading(doc, "14.2 Quy tắc xử lý Admin và Khách", 2)
    add_table(
        doc,
        ["Tình huống", "Quyết định của hệ thống", "Thiết bị phản ứng", "Log hiển thị"],
        [
            ["Admin về nhà", "Cho vào, chạy chế độ Owner/Home", "Mở cửa, bật đèn/quạt theo điều kiện, cho gọi thang máy", "Admin entered - access granted"],
            ["Khách đúng giờ", "Cho vào theo quyền Guest", "Mở cửa, bật đèn lối vào, không bật thiết bị cá nhân", "Guest entered - limited access"],
            ["Khách ngoài giờ", "Từ chối", "Không mở cửa, có thể báo dashboard", "Guest denied - expired schedule"],
            ["Người lạ", "Từ chối và cảnh báo", "Không mở cửa, buzzer, lưu snapshot", "Unknown denied - alert"],
            ["Admin thu hồi quyền khách", "Guest chuyển thành Blocked/Expired", "Lần sau khách không mở được cửa", "Guest permission revoked"],
        ],
        widths=[1.55, 1.85, 1.95, 1.15],
        font_size=7.6,
    )
    add_heading(doc, "14.3 Cách chuẩn bị giấy để demo đẹp", 2)
    add_bullets(
        doc,
        [
            "In mỗi vai trò trên một thẻ A6: Admin màu xanh, Guest màu vàng, Unknown màu đỏ.",
            "Mỗi thẻ nên có ảnh khuôn mặt, tên, vai trò, quyền và thời hạn truy cập.",
            "In một sơ đồ nhà A4 gồm cửa, phòng khách, đèn, quạt, ổ cắm thông minh và thang máy.",
            "Chuẩn bị một tờ log mẫu để đặt cạnh dashboard; khi demo thật thì dashboard sẽ thay tờ log này.",
            "Nếu AI hoặc camera lỗi lúc bảo vệ, vẫn dùng thẻ giấy để kể luồng ra quyết định rồi chuyển sang nút demo thủ công trên dashboard.",
        ],
    )

    add_heading(doc, "15. Hướng phát triển tiếp theo", 1)
    doc.add_paragraph(
        "Sau khi bản mượt hoàn thành, dự án có thể nâng cấp theo nhiều hướng để trở thành một hệ thống smart home gần với sản phẩm thực tế hơn. "
        "Các hướng dưới đây nên làm theo từng giai đoạn, ưu tiên những nâng cấp tăng độ an toàn, độ ổn định và trải nghiệm người dùng."
    )
    add_table(
        doc,
        ["Hướng nâng cấp", "Ý tưởng triển khai", "Giá trị mang lại", "Độ ưu tiên"],
        [
            ["AI chống giả mạo", "Thêm liveness detection: yêu cầu chớp mắt, quay đầu nhẹ, hoặc kiểm tra chuyển động khuôn mặt", "Giảm nguy cơ mở cửa bằng ảnh/video giả", "Rất cao"],
            ["Nhận diện đa yếu tố", "Kết hợp khuôn mặt + RFID/keypad/OTP trong tình huống nhạy cảm", "Tăng bảo mật khi AI không chắc chắn", "Cao"],
            ["Ứng dụng mobile", "App hoặc PWA nhận cảnh báo, mở khóa từ xa, quản lý khách", "Tăng trải nghiệm người dùng, giống sản phẩm thật", "Cao"],
            ["Khách theo lịch", "Tạo mã khách có thời hạn, ví dụ chỉ vào từ 08:00-18:00 trong 3 ngày", "Phù hợp nhà trọ, văn phòng nhỏ, homestay", "Cao"],
            ["Camera snapshot nâng cao", "Lưu ảnh người lạ, ảnh mở cửa thành công, gắn vào access log", "Log có bằng chứng trực quan", "Cao"],
            ["Điều khiển giọng nói offline", "Dùng module nhận giọng nói offline hoặc mô hình speech chạy local", "Demo không phụ thuộc internet", "Trung bình"],
            ["Rule engine thông minh", "Cho phép cấu hình luật: nếu trời tối + Owner về -> bật đèn; nếu vắng nhà -> tắt thiết bị", "Người dùng tự tạo automation", "Cao"],
            ["Cảm biến hiện diện", "Thêm PIR/mmWave để biết phòng còn người hay không", "Tắt thiết bị đúng lúc, tiết kiệm điện", "Trung bình"],
            ["Emergency mode nâng cao", "Gas/flame/khói -> mở cửa, bật quạt hút, gửi cảnh báo, bật đèn thoát hiểm", "Tăng tính an toàn và tính thực tế", "Cao"],
            ["Nguồn dự phòng", "Pin/UPS mini cho khóa và controller cửa", "Mất điện vẫn có thể mở khóa hoặc thoát hiểm", "Cao"],
            ["PCB riêng", "Thiết kế mạch in cho node cửa/smart socket thay vì breadboard", "Mô hình gọn, chuyên nghiệp, ít lỗi dây", "Trung bình"],
            ["Cloud dashboard", "Triển khai backend/web lên cloud, hỗ trợ truy cập từ xa", "Dễ quản lý nhiều nhà/nhiều thiết bị", "Trung bình"],
            ["Phân tích dữ liệu", "Thống kê giờ ra vào, thiết bị dùng nhiều, cảnh báo bất thường", "Biến log thành dữ liệu có ích", "Trung bình"],
            ["Chuẩn hóa sản phẩm", "Đóng vỏ, tài liệu lắp đặt, QR setup Wi-Fi, quy trình reset", "Tiến gần tới prototype thương mại", "Trung bình"],
        ],
        widths=[1.55, 2.25, 2.0, 0.7],
        font_size=7.0,
    )
    add_heading(doc, "15.1 Lộ trình nâng cấp sau bản mượt", 2)
    add_table(
        doc,
        ["Giai đoạn", "Thời gian thêm", "Nên làm", "Kết quả mong đợi"],
        [
            ["Sau bảo vệ 1", "2-4 tuần", "Liveness detection cơ bản, guest theo lịch, Telegram/mobile alert", "Hệ thống an toàn và thuyết phục hơn"],
            ["Sau bảo vệ 2", "1-2 tháng", "PWA/mobile, rule engine, snapshot log nâng cao", "Trải nghiệm giống sản phẩm smart home"],
            ["Prototype nâng cao", "2-3 tháng", "PCB riêng, vỏ hộp, nguồn dự phòng, setup Wi-Fi dễ dùng", "Mô hình gọn và ổn định hơn breadboard"],
            ["Hướng nghiên cứu", "3-6 tháng", "AI edge/local, chống giả mạo tốt hơn, phân tích hành vi bất thường", "Có chiều sâu nghiên cứu AI/IoT"],
        ],
        widths=[1.35, 1.05, 2.75, 1.35],
        font_size=7.6,
    )
    add_callout(
        doc,
        "Ưu tiên sau khi bảo vệ",
        "Nên nâng cấp theo thứ tự: chống giả mạo -> guest theo lịch -> cảnh báo mobile/Telegram -> rule engine -> PCB/vỏ hộp. Đây là đường nâng cấp vừa thực tế, vừa giúp dự án có chiều sâu nếu tiếp tục làm báo cáo hoặc portfolio.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "16. Kết luận triển khai", 1)
    doc.add_paragraph(
        "Để dự án hơn hẳn các mô hình IoT thuần phần cứng, trọng tâm cần đặt vào độ hoàn chỉnh hệ thống: AI nhận diện, phân quyền người dùng, dashboard realtime, log dữ liệu, điều khiển giọng nói, ổ cắm thông minh và mô hình nhà trực quan. "
        "Thứ tự ưu tiên là làm chắc core door trước, sau đó mở rộng sang smart home, voice và thang máy. Nếu gần deadline, cần giữ bản core + smart home hoạt động mượt thay vì cố thêm nhiều cảm biến nhưng thiếu ổn định."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Kế hoạch dự án nhà thông minh AI - Khóa cửa nhận diện khuôn mặt")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
